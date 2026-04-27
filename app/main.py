import hashlib
import json
import math
import os
import re
import shutil
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

import httpx
import numpy as np
from fastapi import FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

DATA_DIR = Path(os.getenv("DATA_DIR", "/app/data"))
UPLOAD_DIR = DATA_DIR / "uploads"
DB_PATH = DATA_DIR / "kb.sqlite3"
APP_PASSWORD = os.getenv("APP_PASSWORD", "changeme")
APP_NAME = os.getenv("APP_NAME", "企业工作资料知识库 AIO")
AUTO_CONFIRM = os.getenv("AUTO_CONFIRM_ON_UPLOAD", "true").lower() == "true"
TOP_K = int(os.getenv("TOP_K", "6"))
OCR_LANG = os.getenv("OCR_LANG", "chi_sim+eng")

LLM_BASE_URL = os.getenv("LLM_BASE_URL", "").rstrip("/")
LLM_API_KEY = os.getenv("LLM_API_KEY", "")
CHAT_MODEL = os.getenv("CHAT_MODEL", "")
EMBEDDING_BASE_URL = os.getenv("EMBEDDING_BASE_URL", "").rstrip("/")
EMBEDDING_API_KEY = os.getenv("EMBEDDING_API_KEY", "")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "")

app = FastAPI(title=APP_NAME)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

STATIC_DIR = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


def ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def db() -> sqlite3.Connection:
    ensure_dirs()
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    with db() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                stored_path TEXT NOT NULL,
                title TEXT,
                doc_type TEXT,
                customer TEXT,
                project TEXT,
                phase TEXT,
                confidentiality TEXT,
                status TEXT,
                tags TEXT,
                summary TEXT,
                created_at TEXT NOT NULL,
                raw_text TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS chunks (
                id TEXT PRIMARY KEY,
                doc_id TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                text TEXT NOT NULL,
                embedding TEXT NOT NULL,
                metadata TEXT,
                FOREIGN KEY(doc_id) REFERENCES documents(id)
            )
            """
        )
        conn.commit()


@app.on_event("startup")
def on_startup() -> None:
    init_db()


def check_password(password: str | None = None, x_app_password: str | None = None) -> None:
    provided = password or x_app_password or ""
    if provided != APP_PASSWORD:
        raise HTTPException(status_code=401, detail="密码错误")


def now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def safe_name(name: str) -> str:
    return re.sub(r"[^\w\-.\u4e00-\u9fff]+", "_", name).strip("_") or "file"


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "gb18030", "gbk", "latin-1"):
        try:
            return raw.decode(enc)
        except Exception:
            pass
    return raw.decode("utf-8", errors="ignore")


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n\n[第{i + 1}页]\n{text}")
        return "\n".join(parts)
    except Exception as exc:
        return f"[PDF解析失败] {exc}"


def extract_docx(path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        return f"[DOCX解析失败] {exc}"


def extract_image(path: Path) -> str:
    try:
        from PIL import Image
        import pytesseract
        return pytesseract.image_to_string(Image.open(path), lang=OCR_LANG)
    except Exception as exc:
        return f"[OCR失败] {exc}"


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in [".txt", ".md", ".csv", ".log"]:
        return read_text_file(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"]:
        return extract_image(path)
    return read_text_file(path)


def contains_any(text: str, words: List[str]) -> bool:
    return any(w in text for w in words)


def detect_doc_type(text: str, filename: str) -> str:
    t = filename + "\n" + text[:6000]
    rules = [
        ("招标文件", ["招标文件", "招标公告", "招标人", "投标截止", "评标办法", "废标"]),
        ("投标文件", ["投标文件", "投标函", "投标人", "投标报价", "投标响应"]),
        ("合同", ["合同", "甲方", "乙方", "合同编号", "违约责任", "付款方式"]),
        ("验收报告", ["验收报告", "初验", "终验", "验收意见", "验收结论"]),
        ("报价单", ["报价单", "报价明细", "报价", "单价", "总价"]),
        ("发票", ["发票", "税率", "购买方", "销售方", "价税合计"]),
        ("项目方案", ["建设方案", "技术方案", "实施方案", "总体架构", "功能模块"]),
        ("会议纪要", ["会议纪要", "会议时间", "参会人员", "会议议题"]),
    ]
    for doc_type, words in rules:
        if contains_any(t, words):
            return doc_type
    return "其他"


def detect_tags(text: str, doc_type: str) -> List[str]:
    tag_rules = {
        "条款/付款": ["付款", "支付", "预付款", "进度款", "尾款", "付款条件"],
        "条款/验收": ["验收", "初验", "终验", "验收标准", "验收条件"],
        "条款/质保": ["质保", "保修", "免费维护", "维保", "缺陷责任期"],
        "条款/违约责任": ["违约", "违约金", "赔偿", "责任承担"],
        "条款/保密": ["保密", "商业秘密", "保密义务"],
        "条款/数据安全": ["数据安全", "网络安全", "个人信息", "安全责任"],
        "条款/知识产权": ["知识产权", "著作权", "源代码", "版权"],
        "阶段/售前": ["需求调研", "方案", "汇报"],
        "阶段/投标": ["招标", "投标", "评标", "废标"],
        "阶段/合同": ["合同", "甲方", "乙方"],
        "阶段/验收": ["验收", "交付", "上线"],
        "风险/付款周期长": ["终验后", "一年后", "分期支付", "尾款"],
        "风险/无限责任": ["全部损失", "无限责任", "连带责任"],
        "状态/已确认": [],
        "状态/可入RAG": [],
    }
    tags = []
    for tag, words in tag_rules.items():
        if words and contains_any(text, words):
            tags.append(tag)
    if doc_type == "合同":
        tags.append("阶段/合同")
    elif doc_type in ["招标文件", "投标文件"]:
        tags.append("阶段/投标")
    elif doc_type in ["验收报告", "交付文档"]:
        tags.append("阶段/验收")
    if AUTO_CONFIRM:
        tags += ["状态/已确认", "状态/可入RAG"]
    return sorted(set(tags))


def regex_first(patterns: List[str], text: str) -> str:
    for p in patterns:
        m = re.search(p, text, re.M)
        if m:
            value = (m.group(1) if m.groups() else m.group(0)).strip()
            return re.sub(r"[：:\s]+$", "", value)[:80]
    return ""


def classify(text: str, filename: str) -> Dict[str, Any]:
    doc_type = detect_doc_type(text, filename)
    tags = detect_tags(text, doc_type)
    customer = regex_first([
        r"甲方[：:\s]*([^\n，,。；;]+)",
        r"招标人[：:\s]*([^\n，,。；;]+)",
        r"采购人[：:\s]*([^\n，,。；;]+)",
        r"客户[：:\s]*([^\n，,。；;]+)",
    ], text)
    project = regex_first([
        r"项目名称[：:\s]*([^\n，,。；;]+)",
        r"工程名称[：:\s]*([^\n，,。；;]+)",
        r"([^\n]{2,40}(?:平台|系统|项目|工程))",
    ], filename + "\n" + text[:2000])
    summary = summarize(text)
    return {
        "title": Path(filename).stem,
        "doc_type": doc_type,
        "customer": customer or "未识别",
        "project": project or "未识别",
        "phase": infer_phase(doc_type),
        "confidentiality": "内部" if doc_type in ["合同", "招标文件", "投标文件"] else "普通",
        "status": "可入RAG" if AUTO_CONFIRM else "待人工确认",
        "tags": tags,
        "summary": summary,
    }


def infer_phase(doc_type: str) -> str:
    if doc_type in ["招标文件", "投标文件"]:
        return "投标"
    if doc_type == "合同":
        return "合同"
    if doc_type in ["验收报告", "交付文档"]:
        return "验收"
    if doc_type in ["项目方案", "报价单"]:
        return "售前"
    return "其他"


def summarize(text: str) -> str:
    lines = [re.sub(r"\s+", " ", x).strip() for x in text.splitlines()]
    lines = [x for x in lines if len(x) >= 8]
    return "\n".join(lines[:5])[:500] if lines else "暂无摘要"


def split_chunks(text: str, max_len: int = 650, overlap: int = 80) -> List[str]:
    text = re.sub(r"\r\n?", "\n", text).strip()
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n{2,}|(?<=。)\s*", text) if p.strip()]
    chunks: List[str] = []
    buf = ""
    for p in paras:
        if len(buf) + len(p) + 1 <= max_len:
            buf = (buf + "\n" + p).strip()
        else:
            if buf:
                chunks.append(buf)
            if len(p) <= max_len:
                buf = p
            else:
                start = 0
                while start < len(p):
                    chunks.append(p[start:start + max_len])
                    start += max_len - overlap
                buf = ""
    if buf:
        chunks.append(buf)
    return chunks[:300]


def tokenize(text: str) -> List[str]:
    text = text.lower()
    zh = re.findall(r"[\u4e00-\u9fff]{1,4}", text)
    en = re.findall(r"[a-z0-9_\-]{2,}", text)
    terms = zh + en
    # Add useful Chinese bigrams/trigrams for semantic-ish matching.
    compact = re.sub(r"\s+", "", text)
    terms += [compact[i:i+2] for i in range(max(0, len(compact)-1)) if re.match(r"[\u4e00-\u9fff]{2}", compact[i:i+2])]
    terms += [compact[i:i+3] for i in range(max(0, len(compact)-2)) if re.match(r"[\u4e00-\u9fff]{3}", compact[i:i+3])]
    return terms[:3000]


def local_embedding(text: str, dim: int = 384) -> List[float]:
    vec = np.zeros(dim, dtype=np.float32)
    for token in tokenize(text):
        h = int(hashlib.sha1(token.encode("utf-8")).hexdigest(), 16)
        idx = h % dim
        sign = 1 if ((h >> 8) & 1) else -1
        vec[idx] += sign
    norm = float(np.linalg.norm(vec))
    if norm:
        vec = vec / norm
    return vec.tolist()


async def remote_embedding(text: str) -> List[float] | None:
    if not (EMBEDDING_BASE_URL and EMBEDDING_API_KEY and EMBEDDING_MODEL):
        return None
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            r = await client.post(
                f"{EMBEDDING_BASE_URL}/embeddings",
                headers={"Authorization": f"Bearer {EMBEDDING_API_KEY}"},
                json={"model": EMBEDDING_MODEL, "input": text[:8000]},
            )
            r.raise_for_status()
            return r.json()["data"][0]["embedding"]
    except Exception:
        return None


def cosine(a: List[float], b: List[float]) -> float:
    if not a or not b:
        return 0.0
    n = min(len(a), len(b))
    aa = a[:n]
    bb = b[:n]
    dot = sum(x * y for x, y in zip(aa, bb))
    na = math.sqrt(sum(x * x for x in aa))
    nb = math.sqrt(sum(y * y for y in bb))
    return dot / (na * nb + 1e-9)


async def embed(text: str) -> List[float]:
    remote = await remote_embedding(text)
    return remote if remote else local_embedding(text)


async def llm_answer(question: str, contexts: List[Dict[str, Any]]) -> str | None:
    if not (LLM_BASE_URL and LLM_API_KEY and CHAT_MODEL):
        return None
    context_text = "\n\n".join([
        f"[来源{i+1}] 文件：{c['filename']}；类型：{c['doc_type']}；片段：{c['text']}"
        for i, c in enumerate(contexts)
    ])
    prompt = f"""你是企业工作资料知识库助手。请只基于给定资料回答问题，不要编造。\n\n问题：{question}\n\n资料：\n{context_text}\n\n要求：\n1. 先给结论。\n2. 列出依据。\n3. 如果资料不足，明确说资料不足。\n"""
    try:
        async with httpx.AsyncClient(timeout=90) as client:
            r = await client.post(
                f"{LLM_BASE_URL}/chat/completions",
                headers={"Authorization": f"Bearer {LLM_API_KEY}"},
                json={
                    "model": CHAT_MODEL,
                    "messages": [
                        {"role": "system", "content": "你是严谨的企业资料问答助手，回答必须可追溯。"},
                        {"role": "user", "content": prompt},
                    ],
                    "temperature": 0.2,
                },
            )
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"]
    except Exception:
        return None


def row_to_doc(row: sqlite3.Row) -> Dict[str, Any]:
    d = dict(row)
    d["tags"] = json.loads(d.get("tags") or "[]")
    d.pop("raw_text", None)
    return d


class AskRequest(BaseModel):
    password: str
    question: str
    top_k: int | None = None


@app.get("/")
def root() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/api/health")
def health() -> Dict[str, Any]:
    return {"ok": True, "app": APP_NAME, "time": now(), "data_dir": str(DATA_DIR)}


@app.post("/api/login")
def login(payload: Dict[str, str]) -> Dict[str, Any]:
    check_password(payload.get("password"))
    return {"ok": True}


@app.get("/api/documents")
def list_documents(password: str | None = None, x_app_password: str | None = Header(default=None)) -> Dict[str, Any]:
    check_password(password, x_app_password)
    with db() as conn:
        rows = conn.execute("SELECT * FROM documents ORDER BY created_at DESC").fetchall()
        return {"documents": [row_to_doc(r) for r in rows]}


@app.post("/api/upload")
async def upload(
    file: UploadFile = File(...),
    password: str = Form(""),
    x_app_password: str | None = Header(default=None),
) -> Dict[str, Any]:
    check_password(password, x_app_password)
    ensure_dirs()
    doc_id = str(uuid.uuid4())
    original = safe_name(file.filename or "upload.bin")
    stored = UPLOAD_DIR / f"{doc_id}_{original}"
    with stored.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    text = extract_text(stored).strip()
    meta = classify(text, original)
    chunks = split_chunks(text)

    with db() as conn:
        conn.execute(
            """
            INSERT INTO documents VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                doc_id,
                original,
                str(stored),
                meta["title"],
                meta["doc_type"],
                meta["customer"],
                meta["project"],
                meta["phase"],
                meta["confidentiality"],
                meta["status"],
                json.dumps(meta["tags"], ensure_ascii=False),
                meta["summary"],
                now(),
                text,
            ),
        )
        for idx, chunk in enumerate(chunks):
            emb = await embed(chunk)
            conn.execute(
                "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?)",
                (str(uuid.uuid4()), doc_id, idx, chunk, json.dumps(emb), json.dumps(meta, ensure_ascii=False)),
            )
        conn.commit()
    return {"ok": True, "document": {"id": doc_id, "filename": original, **meta, "chunks": len(chunks)}}


@app.post("/api/ask")
async def ask(req: AskRequest) -> Dict[str, Any]:
    check_password(req.password)
    q_emb = await embed(req.question)
    limit = req.top_k or TOP_K
    contexts: List[Dict[str, Any]] = []
    with db() as conn:
        rows = conn.execute(
            """
            SELECT c.id as chunk_id, c.text, c.embedding, c.chunk_index,
                   d.id as doc_id, d.filename, d.title, d.doc_type, d.customer, d.project, d.tags
            FROM chunks c JOIN documents d ON c.doc_id = d.id
            """
        ).fetchall()
        for r in rows:
            emb = json.loads(r["embedding"])
            score = cosine(q_emb, emb)
            item = dict(r)
            item["score"] = score
            item["tags"] = json.loads(item.get("tags") or "[]")
            item.pop("embedding", None)
            contexts.append(item)
    contexts.sort(key=lambda x: x["score"], reverse=True)
    contexts = contexts[:limit]
    answer = await llm_answer(req.question, contexts)
    if not answer:
        if not contexts:
            answer = "当前知识库里还没有可检索资料。请先上传合同、标书或方案文件。"
        else:
            parts = ["基于当前检索结果，最相关的依据如下："]
            for i, c in enumerate(contexts[:3], 1):
                snippet = re.sub(r"\s+", " ", c["text"]).strip()[:450]
                parts.append(f"\n{i}. 《{c['filename']}》相关片段：{snippet}")
            answer = "\n".join(parts)
    return {"answer": answer, "sources": contexts}


@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: str, password: str | None = None, x_app_password: str | None = Header(default=None)) -> Dict[str, Any]:
    check_password(password, x_app_password)
    with db() as conn:
        row = conn.execute("SELECT stored_path FROM documents WHERE id=?", (doc_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="文档不存在")
        conn.execute("DELETE FROM chunks WHERE doc_id=?", (doc_id,))
        conn.execute("DELETE FROM documents WHERE id=?", (doc_id,))
        conn.commit()
    try:
        Path(row["stored_path"]).unlink(missing_ok=True)
    except Exception:
        pass
    return {"ok": True}
